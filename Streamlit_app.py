# import streamlit as st
# from docx import Document
# from io import BytesIO

# options = []
# def is_bold(run):
#     return run.font.bold is not None
# def is_highlighted(run):
#     return run.font.highlight_color is not None

# def process_question(paragraph):
#     for run in paragraph.runs:
#         options.append(run.text.strip())
#         if is_highlighted(run) or is_bold(run):
#             return run.text.strip()

#     return None

# def create_answer_document(input_file, output_file):
#     doc = Document(input_file)
#     print(doc)

#     # Process each paragraph and create the answer format
#     answer_doc = Document()
#     current_question = None
#     correct_option = None

#     # Create a BytesIO buffer to store the output document
#     output_buffer = BytesIO()

#     with open(output_file[:-5] + ".txt", "w", encoding="utf-8") as txt_output:
#         for paragraph in doc.paragraphs:
#             text = paragraph.text.strip()

#             if '?' in text:
#                 current_question, options_text1 = text.split('?', 1)
#                 current_question = current_question.strip()
#                 options_text1 = options_text1.strip()
#                 txt_output.write(f"{current_question}\n{options_text1}\n")
#             elif current_question:
#                 answer = process_question(paragraph)
#                 options_text = "\n".join([f"{chr(65 + i)}. {option}" for i, option in enumerate(options)])
#                 print(options_text)
#                 if answer is not None:
#                     correct_option = options.index(answer)
#                     txt_output.write(f"{options_text}\nANSWER: {chr(65 + correct_option)}\n\n")
#                     answer_text = f"{current_question}\n{options_text1}\n{options_text}\nANSWER: {chr(65 + correct_option)}"
#                     answer_doc.add_paragraph(answer_text)
#                     current_question = None
#                     options.clear()
#                     print("Created")

#     # Save the new document to the buffer
#     answer_doc.save(output_buffer)
#     output_buffer.seek(0)  # Reset the buffer position to the beginning

#     return output_buffer.getvalue()  # Return the bytes data

# def main():
#     st.title("Document Processing App")
#     uploaded_file = st.file_uploader("Upload a file", type=["docx"])
#     if uploaded_file:
#         st.write("File uploaded successfully!")
#         output_data = create_answer_document(uploaded_file, "output.docx")
#         st.write("Process Successful")
#         st.download_button(label="Download Output Document", data=output_data, file_name="output.docx")

# if __name__ == "__main__":
#     main()


import streamlit as st
from docx import Document
from io import BytesIO

options = []

def is_bold(run):
    return run.font.bold is not None

def is_highlighted(run):
    return run.font.highlight_color is not None

def process_question(paragraph):
    for run in paragraph.runs:
        options.append(run.text.strip())
        if is_highlighted(run) or is_bold(run):
            return run.text.strip()

    return None

def create_answer_document(input_file, output_file):
    doc = Document(input_file)

    # Process each paragraph and create the answer format
    answer_text = ""
    current_question = None
    options.clear()

    with open(output_file[:-5] + ".txt", "w", encoding="utf-8") as txt_output:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if '?' in text:
                current_question, options_text1 = text.split('?', 1)
                current_question = current_question.strip()
                options_text1 = options_text1.strip()
                txt_output.write(f"{current_question}\n{options_text1}\n")
            elif current_question:
                answer = process_question(paragraph)
                options_text = "\n".join([f"{chr(65 + i)}. {option}" for i, option in enumerate(options)])
                if answer is not None:
                    correct_option = options.index(answer)
                    txt_output.write(f"{options_text}\nANSWER: {chr(65 + correct_option)}\n\n")
                    answer_text += f"{current_question}\n{options_text1}\n{options_text}\nANSWER: {chr(65 + correct_option)}\n\n"
                    current_question = None
                    options.clear()

    # Save the answer text to the output file
    with open(output_file, "w", encoding="utf-8") as output_file:
        output_file.write(answer_text)

    return answer_text

def main():
    st.title("Document Processing App")
    st.write("Fomat of Every Question Should be like this: \n A “Payee’s Account Only” Crossed cheque can be deposited in account having title different from Payee name mentioned in Cheque? \n A.	True \n B. False \n")
    uploaded_file = st.file_uploader("Upload a file", type=["docx"])
    if uploaded_file:
        # st.write("File uploaded successfully!")
        output_data = create_answer_document(uploaded_file, "output.docx")
        # st.write("Process Successful")
        
        # Display the output text in a scrollable area
        st.text_area("Output Text", output_data, height=400)

        # Render download buttons inline
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(label="Download Output Document (DOCX)", data=output_data, file_name="output.docx")
        with col2:
            st.download_button(label="Download Output Document (TXT)", data=output_data.encode("utf-8"), file_name="output.txt")

if __name__ == "__main__":
    main()

