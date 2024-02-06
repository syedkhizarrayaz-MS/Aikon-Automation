from docx import Document

options = []

def is_highlighted(run):
    return run.font.highlight_color is not None

def process_question(paragraph):
    for run in paragraph.runs:
        options.append(run.text.strip())
        if is_highlighted(run):
            return run.text.strip()

    return None

def create_answer_document(input_file, output_file):
    doc = Document(input_file)
    
    # Process each paragraph and create the answer format
    answer_doc = Document()
    current_question = None
    correct_option = None

    # Create a text file for output
    with open(output_file[:-5] + ".txt", "w", encoding="utf-8") as txt_output:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if '?' in text:
                current_question, options_text1 = text.split('?', 1)
                current_question = current_question.strip()
                options_text1 = options_text1.strip()
                txt_output.write(f"{current_question}\n{options_text1}\n")
            elif current_question:
                answer = process_question(paragraph)
                options_text = "\n".join([f"{chr(65 + i)}. {option}" for i, option in enumerate(options)])
                if answer is not None:
                    correct_option = options.index(answer)
                    txt_output.write(f"{options_text}\nANSWER: {chr(65 + correct_option)}\n\n")
                    answer_text = f"{current_question}\n{options_text1}\n{options_text}\nANSWER: {chr(65 + correct_option)}"
                    print(f"Answer found: {answer_text}")
                    answer_doc.add_paragraph(answer_text)
                    current_question = None
                    options.clear()

    # Save the new document
    answer_doc.save(output_file)

if __name__ == "__main__":
    input_file = r"C:\Users\Khizar Riyaz\Downloads\Product Knowledge - Liability Products (100 QUES).docx"
    output_file = r"C:\Users\Khizar Riyaz\Downloads\output.docx"

    # Create the answer document
    create_answer_document(input_file, output_file)

    print(f"\nAnswer document created: {output_file}")
